import streamlit as st
import docx
from faker import Faker
import pycountry
from deep_translator import GoogleTranslator
import spacy
import nltk
from nltk.tokenize import sent_tokenize
from iso639 import languages
from faker.config import AVAILABLE_LOCALES


def main():
    nltk.download('punkt')
    spacy.cli.download("en_core_web_sm")
    language_model = "en_core_web_sm"
    nlp = spacy.load(language_model)
    st.title("Document Translation App")

    # File uploader widget
    uploaded_file = st.file_uploader("Upload a docx file", type=["docx"])

    target_language = st.selectbox("Select target language:", [lang.name for lang in languages])
    target_country = st.selectbox("Select target country:", [country.name for country in pycountry.countries])

    if uploaded_file and st.button("Localize"):
            translated_doc = translate_and_generate_output(nlp, uploaded_file, target_language, target_country)
            download_link = generate_download_link(translated_doc, target_language)
            st.markdown(download_link, unsafe_allow_html=True)
    else:
            st.error("Translation failed. Please check your input and try again.")


def generate_unique_names(nlp, text, country_code, name_mapping):
    doc = nlp(text)
    unique_names = set()

    for ent in doc.ents:
        if ent.label_ == "PERSON":
            unique_names.add(ent.text)

    faker = Faker(country_code)

    for name in unique_names:
      if name not in name_mapping.keys():
        name_mapping[name] = faker.name()

    return name_mapping

def get_country_code(country_name):
    try:
        country = pycountry.countries.search_fuzzy(country_name)
        return country[0].alpha_2
    except AttributeError:
        return 'IN'


def get_short_form(language_name):
    for lang in languages:
        if lang.name.lower() == language_name.lower():
            return lang.alpha2
    return 'en'

def translate_text_2(text, language):
  sentences = sent_tokenize(text)
  translated_list = []
  for sentence in sentences:
    translated = GoogleTranslator(source='auto', target=language).translate(sentence)
    translated_list.append(translated)

  return "".join(translated_list)

def translate_and_generate_output(nlp, uploaded_file, target_language, target_country):
    # Read the uploaded file and convert to docx.Document
    doc = docx.Document(uploaded_file)
    out_doc = docx.Document()
    faker_code = "en_US"
    name_mapping = {}
    country = get_country_code(target_country)
    language = get_short_form(target_language)
    for locale in AVAILABLE_LOCALES:
        if country in locale :
            faker_code = locale
            break
    print("mapping")
    for para in doc.paragraphs:
        if para.text.strip() == "":
            continue
        else:
            name_mapping = generate_unique_names(nlp, para.text, faker_code, name_mapping)

    for para in doc.paragraphs:
        if para.text.strip() == "":
            out_doc.add_paragraph().paragraph_format.line_spacing = 0.75
        else:
            for original_name, new_name in name_mapping.items():
                para.text = para.text.replace(original_name, new_name)
            out_doc.add_paragraph(para.text).paragraph_format.line_spacing = 0.75

    translated_doc = docx.Document()
    for para in out_doc.paragraphs:
        if para.text.strip() == "":
            translated_doc.add_paragraph().paragraph_format.line_spacing = 1.15
        else:
            translated_text = translate_text_2(para.text, language)
            translated_doc.add_paragraph(translated_text).paragraph_format.line_spacing = 1.15

    return translated_doc

def generate_download_link(translated_doc, target_language):
    # Save the translated_doc to a temporary file
    temp_file_path = "{target_language}_translated_doc.docx"
    translated_doc.save(temp_file_path)

    # Generate a download link for the user
    download_link = f'<a href="{temp_file_path}" download="{target_language}_translated.docx">Click here to download translated document</a>'
    return download_link


if __name__ == "__main__":

    main()
